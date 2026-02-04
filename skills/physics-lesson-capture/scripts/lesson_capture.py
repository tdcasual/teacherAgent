#!/usr/bin/env python3
import argparse
import json
import shutil
import sys
from pathlib import Path
from typing import List, Tuple

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from ocr_utils import load_env_from_dotenv, ocr_with_sdk


def detect_type(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}:
        return "image"
    if ext in {".pdf"}:
        return "pdf"
    if ext in {".docx"}:
        return "docx"
    return "unknown"


def extract_pdf_text(path: Path) -> List[Tuple[int, str]]:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        return []
    pages_text = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages_text.append((idx, text))
    except Exception:
        return []
    return pages_text


def extract_docx_text(path: Path) -> str:
    try:
        import docx  # type: ignore
    except Exception:
        return ""
    try:
        doc = docx.Document(str(path))
        lines = [p.text for p in doc.paragraphs if p.text]
        return "\n".join(lines)
    except Exception:
        return ""


def extract_docx_images(path: Path, out_dir: Path) -> List[Path]:
    try:
        import docx  # type: ignore
    except Exception:
        return []
    images = []
    try:
        doc = docx.Document(str(path))
        for rel in doc.part._rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                img_name = Path(rel.target_ref).name
                dest = out_dir / f"{path.stem}_{img_name}"
                dest.write_bytes(img_data)
                images.append(dest)
    except Exception:
        return []
    return images


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def extract_examples_from_text(text: str) -> List[dict]:
    examples = []
    lines = text.splitlines()

    def is_q_start(line: str) -> bool:
        line = line.strip()
        if not line:
            return False
        return bool(__import__("re").match(r"^\d+[\.、]\s*", line))

    def is_option(line: str) -> bool:
        line = line.strip()
        return bool(__import__("re").match(r"^[A-D][\.、]\s*", line))

    current = {"stem": "", "options": []}
    for line in lines:
        if is_q_start(line):
            if current["stem"]:
                examples.append(current)
            current = {"stem": line.strip(), "options": []}
            continue
        if is_option(line):
            current["options"].append(line.strip())
            continue
        if current["stem"]:
            # append to stem if in the same block
            current["stem"] += " " + line.strip()

    if current["stem"]:
        examples.append(current)

    return examples


def main():
    parser = argparse.ArgumentParser(description="Capture lesson materials, OCR, and extract examples")
    parser.add_argument("--lesson-id", required=True)
    parser.add_argument("--topic", required=True)
    parser.add_argument("--class-name", default="")
    parser.add_argument("--sources", nargs="+", required=True, help="PDF/DOCX/Image files")
    parser.add_argument("--discussion-notes", help="teacher discussion notes (md)")
    parser.add_argument("--lesson-plan", help="lesson plan (md/docx/pdf)")
    parser.add_argument("--force-ocr", action="store_true", help="force OCR even if PDF has text")
    parser.add_argument("--ocr-mode", default="FREE_OCR", help="OCR mode for DeepSeek-OCR")
    parser.add_argument("--language", default="zh")
    parser.add_argument("--out-base", default="data/lessons")
    args = parser.parse_args()

    load_env_from_dotenv(Path('.env'))

    lesson_dir = Path(args.out_base) / args.lesson_id
    sources_dir = lesson_dir / "sources"
    ocr_dir = lesson_dir / "ocr"
    text_dir = lesson_dir / "text"
    examples_dir = lesson_dir / "examples"
    sources_dir.mkdir(parents=True, exist_ok=True)
    ocr_dir.mkdir(parents=True, exist_ok=True)
    text_dir.mkdir(parents=True, exist_ok=True)
    examples_dir.mkdir(parents=True, exist_ok=True)

    manifest = {
        "lesson_id": args.lesson_id,
        "topic": args.topic,
        "class_name": args.class_name,
        "sources": [],
    }

    # copy sources
    source_paths = []
    for src in args.sources:
        src_path = Path(src)
        if not src_path.exists():
            raise SystemExit(f"Source not found: {src}")
        dest = sources_dir / src_path.name
        if dest.resolve() != src_path.resolve():
            shutil.copy2(str(src_path), str(dest))
        source_paths.append(dest)
        manifest["sources"].append({"file": dest.name, "type": detect_type(dest)})

    # include lesson plan if provided
    if args.lesson_plan:
        lp = Path(args.lesson_plan)
        if lp.exists():
            dest = sources_dir / lp.name
            if dest.resolve() != lp.resolve():
                shutil.copy2(str(lp), str(dest))
            manifest["sources"].append({"file": dest.name, "type": detect_type(dest)})
            source_paths.append(dest)

    # write manifest
    (lesson_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    # extract text + OCR
    collected_texts = []
    for src in source_paths:
        ftype = detect_type(src)
        if ftype == "pdf":
            pages_text = extract_pdf_text(src)
            text_concat = "\n".join([t for _, t in pages_text if t])
            if args.force_ocr or len(text_concat.strip()) < 50:
                # OCR whole pdf
                try:
                    ocr_text = ocr_with_sdk(src, mode=args.ocr_mode, language=args.language)
                except Exception as exc:
                    raise SystemExit(f"OCR failed for {src}: {exc}")
                (ocr_dir / f"{src.stem}.json").write_text(json.dumps({"text": ocr_text}, ensure_ascii=False, indent=2), encoding="utf-8")
                cleaned = normalize_text(ocr_text)
                (text_dir / f"{src.stem}.txt").write_text(cleaned, encoding="utf-8")
                collected_texts.append((src.name, cleaned))
            else:
                # save per page
                combined = []
                for page_no, page_text in pages_text:
                    cleaned = normalize_text(page_text)
                    (text_dir / f"{src.stem}_p{page_no}.txt").write_text(cleaned, encoding="utf-8")
                    collected_texts.append((f"{src.name}#p{page_no}", cleaned))
                    combined.append(cleaned)
                if combined:
                    (text_dir / f"{src.stem}.txt").write_text("\n".join(combined), encoding="utf-8")
        elif ftype == "docx":
            text = extract_docx_text(src)
            if text:
                cleaned = normalize_text(text)
                (text_dir / f"{src.stem}.txt").write_text(cleaned, encoding="utf-8")
                collected_texts.append((src.name, cleaned))
            # OCR embedded images if any
            img_dir = lesson_dir / "sources" / "docx_images"
            img_dir.mkdir(parents=True, exist_ok=True)
            images = extract_docx_images(src, img_dir)
            for img in images:
                try:
                    ocr_text = ocr_with_sdk(img, mode=args.ocr_mode, language=args.language)
                except Exception as exc:
                    raise SystemExit(f"OCR failed for {img}: {exc}")
                (ocr_dir / f"{img.stem}.json").write_text(json.dumps({"text": ocr_text}, ensure_ascii=False, indent=2), encoding="utf-8")
                cleaned = normalize_text(ocr_text)
                (text_dir / f"{img.stem}.txt").write_text(cleaned, encoding="utf-8")
                collected_texts.append((img.name, cleaned))
            if not text and not images:
                # fallback to OCR on whole docx (some SDKs accept docx)
                try:
                    ocr_text = ocr_with_sdk(src, mode=args.ocr_mode, language=args.language)
                except Exception as exc:
                    raise SystemExit(f"OCR failed for {src}: {exc}")
                (ocr_dir / f"{src.stem}.json").write_text(json.dumps({"text": ocr_text}, ensure_ascii=False, indent=2), encoding="utf-8")
                cleaned = normalize_text(ocr_text)
                (text_dir / f"{src.stem}.txt").write_text(cleaned, encoding="utf-8")
                collected_texts.append((src.name, cleaned))
        elif ftype == "image":
            try:
                ocr_text = ocr_with_sdk(src, mode=args.ocr_mode, language=args.language)
            except Exception as exc:
                raise SystemExit(f"OCR failed for {src}: {exc}")
            (ocr_dir / f"{src.stem}.json").write_text(json.dumps({"text": ocr_text}, ensure_ascii=False, indent=2), encoding="utf-8")
            cleaned = normalize_text(ocr_text)
            (text_dir / f"{src.stem}.txt").write_text(cleaned, encoding="utf-8")
            collected_texts.append((src.name, cleaned))
        else:
            # unknown: try OCR
            try:
                ocr_text = ocr_with_sdk(src, mode=args.ocr_mode, language=args.language)
            except Exception as exc:
                raise SystemExit(f"OCR failed for {src}: {exc}")
            (ocr_dir / f"{src.stem}.json").write_text(json.dumps({"text": ocr_text}, ensure_ascii=False, indent=2), encoding="utf-8")
            cleaned = normalize_text(ocr_text)
            (text_dir / f"{src.stem}.txt").write_text(cleaned, encoding="utf-8")
            collected_texts.append((src.name, cleaned))

    # extract examples
    examples = []
    counter = 1
    for source_ref, text in collected_texts:
        for ex in extract_examples_from_text(text):
            example_id = f"E{counter:03d}"
            counter += 1
            options = " | ".join(ex.get("options", []))
            examples.append({
                "example_id": example_id,
                "lesson_id": args.lesson_id,
                "stem_text": ex.get("stem", ""),
                "options": options,
                "answer": "",
                "kp_candidate": "",
                "difficulty": "",
                "source_ref": source_ref,
                "notes": "",
            })

    # write examples.csv
    if examples:
        import csv
        csv_path = lesson_dir / "examples.csv"
        with csv_path.open("w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=list(examples[0].keys()))
            writer.writeheader()
            for row in examples:
                writer.writerow(row)

    # class discussion
    discussion_path = lesson_dir / "class_discussion.md"
    if args.discussion_notes:
        src = Path(args.discussion_notes)
        if src.exists():
            discussion_path.write_text(src.read_text(encoding="utf-8"), encoding="utf-8")
    else:
        discussion_path.write_text(
            f"Lesson: {args.topic}\nKey Misconceptions:\n- \nEmphasis KP:\n- \nClass Observations:\n- \nHomework Focus:\n- ",
            encoding="utf-8",
        )

    # lesson summary
    summary_path = lesson_dir / "lesson_summary.md"
    summary_lines = [
        f"Lesson: {args.topic}",
        f"Lesson ID: {args.lesson_id}",
        "",
        f"Sources: {', '.join([p.name for p in source_paths])}",
        f"Examples extracted: {len(examples)}",
        "",
        "Example Preview:",
    ]
    for ex in examples[:5]:
        preview = ex["stem_text"][:120]
        summary_lines.append(f"- {ex['example_id']}: {preview}")
    summary_path.write_text("\n".join(summary_lines), encoding="utf-8")

    print(f"[OK] Lesson captured at {lesson_dir}")


if __name__ == "__main__":
    main()
